from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "admin" / "Senior-Project-Admin-Overview.docx"

# Resolved preset: standard_business_brief.
# Named overrides used throughout: Aptos typography to match the app, a compact
# 10.5 pt body for the two-page handout, and the app's navy/blue/teal palette.
NAVY = "0B2341"
BLUE = "1769AA"
TEAL = "109C98"
INK = "1C2B3A"
MUTED = "5D6B78"
LINE = "D7E2EA"
PALE_BLUE = "EDF5FB"
PALE_TEAL = "EAF7F5"
PALE_AMBER = "FFF4DC"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF6EF"
WHITE = "FFFFFF"
SOFT_GRAY = "F5F7F9"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_START_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 92, bottom: int = 92, start: int = CELL_START_DXA, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge in edges.items():
        tag = "start" if edge_name == "left" else "end" if edge_name == "right" else edge_name
        element = borders.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            borders.append(element)
        element.set(qn("w:val"), edge.get("val", "single"))
        element.set(qn("w:sz"), str(edge.get("sz", 6)))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), edge.get("color", LINE))


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {PAGE_WIDTH_DXA} DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        old_grid.append(col)

    for row in table.rows:
        keep_row_together(row)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size: float = 10.5, color: str = INK, bold: bool = False, italic: bool = False, name: str = "Aptos") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, *, before: float = 0, after: float = 5, line: float = 1.1, keep_next: bool = False, keep_together: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_together


def clear_cell(cell) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    format_paragraph(p, after=0)


def add_cell_text(cell, label: str, detail: str, *, label_color: str = NAVY, detail_color: str = INK, size: float = 9.4) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    format_paragraph(p, after=0, line=1.08)
    r = p.add_run(label)
    set_font(r, size=size, color=label_color, bold=True)
    r2 = p.add_run(detail)
    set_font(r2, size=size, color=detail_color)


def set_section_header(section) -> None:
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(p, after=0)
    left = p.add_run("SENIOR PROJECT WORKSPACE")
    set_font(left, size=8.5, color=MUTED, bold=True)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5))
    right = p.add_run("\tADMINISTRATOR REFERENCE")
    set_font(right, size=8.5, color=MUTED, bold=True)


def set_section_footer(section, page_number: int) -> None:
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    format_paragraph(p, after=0)
    label = p.add_run(f"Updated September 2026  |  Page {page_number} of 2")
    set_font(label, size=8, color=MUTED)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=2, keep_next=True)
    r = p.add_run(text.upper())
    set_font(r, size=9, color=TEAL, bold=True)


def add_title(doc: Document, text: str, size: float = 27) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=3, line=1.0, keep_next=True)
    r = p.add_run(text)
    set_font(r, size=size, color=NAVY, bold=True)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=11, line=1.1, keep_next=True)
    r = p.add_run(text)
    set_font(r, size=11.5, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), TEAL)
    border.append(bottom)
    p_pr.append(border)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def add_callout(doc: Document, label: str, text: str, *, fill: str = PALE_BLUE, accent: str = BLUE) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, before=8, after=8, line=1.1, keep_together=True)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label}  ")
    set_font(r, size=10.2, color=NAVY, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=10.2, color=INK)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, after=4, line=1.08, keep_together=True)
    r = p.add_run(f"{label}. ")
    set_font(r, size=10.2, color=NAVY, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=10.2, color=INK)


def style_header_row(row, labels: list[str], fill: str = NAVY) -> None:
    set_repeat_table_header(row)
    for cell, label in zip(row.cells, labels):
        shade_cell(cell, fill)
        clear_cell(cell)
        p = cell.paragraphs[0]
        format_paragraph(p, after=0)
        r = p.add_run(label.upper())
        set_font(r, size=8.7, color=WHITE, bold=True)
        set_cell_border(cell, top={"color": fill}, bottom={"color": fill}, left={"color": fill}, right={"color": fill})


def add_data_boundary_table(doc: Document) -> None:
    rows = [
        ("Names, school email, school, program, year, and assigned role", "Google account passwords or Google sign-in details"),
        ("Project name, 1-5 student team, Mentor, and Program Teacher", "Copies of student Docs, Slides, Sheets, photos, or other files"),
        ("Guided answers, personal reflections, drafts, and version history", "The contents or file list inside a student's Drive folder"),
        ("Project-folder, work, and school-template links", "Permission to edit, move, delete, or copy items in Google Drive"),
        ("Statuses, deadlines, review notes, presentation details, and activity records", "Readable passwords or reusable setup codes"),
    ]
    table = doc.add_table(rows=1, cols=2)
    style_header_row(table.rows[0], ["The app keeps", "The app does not keep"])
    for index, (keeps, does_not) in enumerate(rows):
        row = table.add_row()
        fill = WHITE if index % 2 == 0 else SOFT_GRAY
        for cell in row.cells:
            shade_cell(cell, fill)
            set_cell_border(cell, top={"color": LINE, "sz": 4}, bottom={"color": LINE, "sz": 4}, left={"color": LINE, "sz": 4}, right={"color": LINE, "sz": 4})
        add_cell_text(row.cells[0], "", keeps, size=9.25)
        add_cell_text(row.cells[1], "", does_not, size=9.25)
    set_table_geometry(table, [4680, 4680])


def add_access_line(doc: Document) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, before=5, after=0, line=1.08, keep_together=True)
    parts = [
        ("Student", "own project"),
        ("Mentor", "assigned projects"),
        ("Program Teacher", "program projects and reviews"),
        ("School/Site Admin", "local setup and oversight"),
        ("Global Admin", "all schools"),
        ("Viewer", "approved records, read-only"),
    ]
    lead = p.add_run("Access by role  ")
    set_font(lead, size=9.4, color=TEAL, bold=True)
    for index, (role, detail) in enumerate(parts):
        if index:
            sep = p.add_run("  |  ")
            set_font(sep, size=9, color=LINE)
        r = p.add_run(f"{role}: ")
        set_font(r, size=9, color=NAVY, bold=True)
        r2 = p.add_run(detail)
        set_font(r2, size=9, color=INK)


def add_status_table(doc: Document) -> None:
    rows = [
        ("People needed", "A Mentor or Program Teacher has not accepted yet.", "Student or admin tags the person; the adult accepts.", PALE_AMBER, "8C5D00"),
        ("In progress", "The team is working and no item is waiting on staff.", "Student continues the next listed item.", PALE_BLUE, BLUE),
        ("Waiting for review", "At least one item has been turned in.", "Assigned reviewer opens it and records a decision.", PALE_TEAL, TEAL),
        ("Changes needed", "A reviewer returned an item with a clear correction.", "Student revises it and turns it in again.", PALE_RED, "A33A3A"),
        ("Completed / archived", "Active work is finished, or the record is kept for history.", "Admin confirms closeout and keeps the record out of active work.", PALE_GREEN, "2B7A4B"),
    ]
    table = doc.add_table(rows=1, cols=3)
    style_header_row(table.rows[0], ["Project status", "What it means", "What happens next"])
    for status, meaning, next_step, fill, accent in rows:
        row = table.add_row()
        for cell in row.cells:
            shade_cell(cell, WHITE)
            set_cell_border(cell, top={"color": LINE, "sz": 4}, bottom={"color": LINE, "sz": 4}, left={"color": LINE, "sz": 4}, right={"color": LINE, "sz": 4})
        shade_cell(row.cells[0], fill)
        add_cell_text(row.cells[0], "", status, label_color=accent, detail_color=accent, size=9.2)
        row.cells[0].paragraphs[0].runs[0].bold = True
        add_cell_text(row.cells[1], "", meaning, size=9.05)
        add_cell_text(row.cells[2], "", next_step, size=9.05)
    set_table_geometry(table, [2100, 3370, 3890])


def add_flow_table(doc: Document) -> None:
    rows = [
        ("1", "Start + team", "A project is created, or a student submits a new idea and tags teammates. A project may have 1-5 students."),
        ("2", "People", "The Mentor and Program Teacher are tagged. Both must accept before the project is fully ready."),
        ("3", "Proposal", "Students answer short prompts, define the result, and submit the proposal. The Program Teacher approves it before build work moves forward."),
        ("4", "Build I", "Students work in Google Drive, keep a short journal, add the matching links, and prepare for Mentor feedback."),
        ("5", "Build II", "The team updates the work, records how Mentor feedback was used, and prepares the presentation outline and time."),
        ("6", "Present", "Students practice, present the project, show the linked work, and complete school check-in or check-out steps."),
        ("7", "Celebrate + reflect", "The team shares the result. Each student completes personal reflections, thank-you work, and a next-step plan."),
        ("8", "Finish", "Students save final copies somewhere they can keep. Admins confirm closeout and move finished records out of active work."),
    ]
    table = doc.add_table(rows=1, cols=3)
    style_header_row(table.rows[0], ["", "Stage", "What happens"])
    for index, (number, stage, detail) in enumerate(rows):
        row = table.add_row()
        fill = WHITE if index % 2 == 0 else SOFT_GRAY
        for cell in row.cells:
            shade_cell(cell, fill)
            set_cell_border(cell, top={"color": LINE, "sz": 3}, bottom={"color": LINE, "sz": 3}, left={"color": LINE, "sz": 3}, right={"color": LINE, "sz": 3})
        shade_cell(row.cells[0], PALE_BLUE)
        clear_cell(row.cells[0])
        p_num = row.cells[0].paragraphs[0]
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_num = p_num.add_run(number)
        set_font(r_num, size=10.5, color=BLUE, bold=True)
        add_cell_text(row.cells[1], "", stage, detail_color=NAVY, size=9.15)
        row.cells[1].paragraphs[0].runs[0].bold = True
        add_cell_text(row.cells[2], "", detail, size=8.85)
    set_table_geometry(table, [620, 1840, 6900])


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    heading_specs = {
        1: (15.5, NAVY, 10, 5),
        2: (12.5, BLUE, 7, 4),
        3: (11.5, NAVY, 5, 3),
    }
    for level, (size, color, before, after) in heading_specs.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    configure_section(section)

    configure_styles(doc)

    doc.core_properties.title = "Senior Project App - Admin Overview"
    doc.core_properties.subject = "Purpose, data boundaries, project statuses, and project flow"
    doc.core_properties.author = "Senior Project Workspace"
    doc.core_properties.keywords = "senior project, capstone, administrator, overview"

    add_kicker(doc, "Admin overview")
    add_title(doc, "Senior Project Workspace")
    add_subtitle(doc, "What the app does, what it keeps, and how a project moves")

    add_callout(
        doc,
        "In plain terms",
        "This is the school's project organizer and review record. Students write guided responses here and link to work that stays in their own Google Drive.",
    )

    add_heading(doc, "What the app is", 1)
    add_labeled_paragraph(doc, "One project workspace", "Each student belongs to one active project. Projects may be individual or shared by a team of up to five students.")
    add_labeled_paragraph(doc, "A guided process", "Short prompts, journals, reflections, templates, due dates, and one clear next step help students move through the year. Students may draft future work without skipping approval gates.")
    add_labeled_paragraph(doc, "A review loop", "Students turn in an answer or Google Drive link. Mentors and Program Teachers review the work, approve it, or return it with a specific change to make.")
    add_labeled_paragraph(doc, "One shell, different tools", "Everyone opens the same project-centered workspace. Students, Mentors, Program Teachers, viewers, and admins see only the actions their role permits.")

    add_heading(doc, "The data boundary", 1)
    add_data_boundary_table(doc)

    add_access_line(doc)

    page_two = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(page_two)

    add_kicker(doc, "Project operations")
    add_title(doc, "How a project moves", size=25)
    add_subtitle(doc, "The status tells an admin who owns the next action")

    add_heading(doc, "Project status at a glance", 1)
    add_status_table(doc)

    add_callout(
        doc,
        "Work-item loop",
        "Draft -> Turned in -> Approved, or Needs changes -> revise -> turn in again. A project status is a roll-up of the people and work inside it; staff do not have to guess or update it by hand.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_heading(doc, "The project flow", 1)
    add_flow_table(doc)

    p = doc.add_paragraph()
    format_paragraph(p, before=6, after=0, line=1.08, keep_together=True)
    lead = p.add_run("Who moves the work  ")
    set_font(lead, size=9.3, color=TEAL, bold=True)
    roles = [
        ("Student", "creates and revises"),
        ("Mentor", "advises and responds"),
        ("Program Teacher", "reviews and approves"),
        ("Admin", "manages access, projects, templates, and risk"),
    ]
    for index, (role, detail) in enumerate(roles):
        if index:
            sep = p.add_run("  |  ")
            set_font(sep, size=8.9, color=LINE)
        rr = p.add_run(f"{role}: ")
        set_font(rr, size=8.9, color=NAVY, bold=True)
        rd = p.add_run(detail)
        set_font(rd, size=8.9, color=INK)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
